import streamlit as st
import pandas as pd
from datetime import datetime
from io import BytesIO
import xmlrpc.client
import base64
import os
import json
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# ─────────────────────────────────────────────────────────────
# 📁 GOOGLE DRIVE UPLOAD
# ─────────────────────────────────────────────────────────────
GDRIVE_WEBHOOK_URL = os.environ.get("GDRIVE_WEBHOOK_URL", "")


def upload_to_gdrive(filename, file_bytes, mimetype, subfolder_name=None):
    """Upload a file to Google Drive via Apps Script webhook."""
    try:
        import urllib.request

        if not GDRIVE_WEBHOOK_URL:
            return None, "Google Drive not configured (missing GDRIVE_WEBHOOK_URL)"

        payload = json.dumps({
            "filename": filename,
            "data": base64.b64encode(file_bytes).decode("utf-8"),
            "mimetype": mimetype,
            "subfolder": subfolder_name or ""
        }).encode("utf-8")

        req = urllib.request.Request(
            GDRIVE_WEBHOOK_URL,
            data=payload,
            headers={"Content-Type": "application/json"},
            method="POST"
        )
        resp = urllib.request.urlopen(req, timeout=30)
        result = json.loads(resp.read().decode("utf-8"))

        if result.get("success"):
            return result, None
        else:
            return None, result.get("error", "Unknown error")

    except Exception as e:
        return None, str(e)

# Conversion constants
inches_to_mm = 25.4
mm_to_inches = 1 / inches_to_mm
sq_inches_to_sq_feet = 1 / 144


# ─────────────────────────────────────────────────────────────
# 📄 GENERATE PURCHASE ORDER DOCX
# ─────────────────────────────────────────────────────────────
def _add_shading(cell, color):
    """Add background shading to a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = tcPr.makeelement(qn('w:shd'), {qn('w:fill'): color, qn('w:val'): 'clear'})
    tcPr.append(shd)


def generate_po_docx(
    vendor_name, vendor_contact, vendor_email, vendor_address,
    ship_to_lines, job_number, job_location,
    po_date, po_number, requisitioner,
    lead_time, shipped_via, fob_point, terms,
    glass_lines,
    price_per_sqft, packaging_cost, shipping_cost, sales_tax, other_cost,
    packaging_note="Non-returnable boxed crate/rack",
    logo_path=None,
):
    """Generate an INOVUES-format Purchase Order as a .docx BytesIO buffer."""
    doc = DocxDocument()

    for section in doc.sections:
        section.top_margin = Cm(1.27)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)

    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(9)

    subtotal = sum(line['area_total'] * price_per_sqft for line in glass_lines)
    total = subtotal + (sales_tax or 0) + (packaging_cost or 0) + (shipping_cost or 0) + (other_cost or 0)

    table = doc.add_table(rows=0, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Set column widths: ITEM#, DESC, UNIT SIZE, AREA EACH, QTY, TOTAL AREA, TOTAL
    col_widths = [Cm(1.2), Cm(3.0), Cm(3.8), Cm(2.2), Cm(1.2), Cm(2.5), Cm(2.5)]
    for i, width in enumerate(col_widths):
        table.columns[i].width = width

    # ── Row: Logo + PURCHASE ORDER ──
    row = table.add_row()
    c = row.cells[0]; c.merge(row.cells[2])
    if logo_path:
        c.paragraphs[0].add_run().add_picture(logo_path, width=Inches(1.5))
    c = row.cells[3]; c.merge(row.cells[6])
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("PURCHASE ORDER")
    run.bold = True; run.font.size = Pt(16); run.font.small_caps = True

    # ── Row: TO / SHIP TO / BILL TO ──
    row = table.add_row()

    c = row.cells[0]; c.merge(row.cells[1])
    p = c.paragraphs[0]; run = p.add_run("TO:"); run.bold = True; run.font.small_caps = True
    c.add_paragraph().add_run(vendor_name)
    if vendor_contact:
        c.add_paragraph().add_run(f"Attn: {vendor_contact}")
    if vendor_email:
        c.add_paragraph().add_run(vendor_email)
    if vendor_address:
        for ln in vendor_address.split('\n'):
            if ln.strip():
                c.add_paragraph().add_run(ln)
    c.add_paragraph()
    p = c.add_paragraph(); run = p.add_run(f"JOB NO.: {job_number}"); run.bold = True
    if job_location:
        p = c.add_paragraph(); run = p.add_run(f"JOB LOCATION: {job_location}"); run.bold = True

    c = row.cells[2]; c.merge(row.cells[4])
    p = c.paragraphs[0]; run = p.add_run("SHIP TO:"); run.bold = True; run.font.small_caps = True
    for ln in ship_to_lines:
        c.add_paragraph().add_run(ln)

    c = row.cells[5]; c.merge(row.cells[6])
    p = c.paragraphs[0]; run = p.add_run("BILL TO:"); run.bold = True; run.font.small_caps = True
    for ln in ["INOVUES, INC.", "2700 Post Oak Blvd., 2100", "Houston, TX 77056",
               "accounts@inovues.com", "(833) 466-8837 (INO-VUES)"]:
        c.add_paragraph().add_run(ln)

    # ── Spacer ──
    row = table.add_row(); row.cells[0].merge(row.cells[6])

    # ── PO metadata headers ──
    row = table.add_row()
    for i, h in enumerate(["P.O. DATE", "P.O. NUMBER", "REQUISITIONER", "LEAD TIME",
                            "SHIPPED VIA", "F.O.B. POINT", "TERMS"]):
        p = row.cells[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h); run.bold = True; run.font.size = Pt(7); run.font.small_caps = True
        _add_shading(row.cells[i], 'D9E2F3')

    # ── PO metadata values ──
    row = table.add_row()
    for i, v in enumerate([po_date, po_number, requisitioner, lead_time, shipped_via, fob_point, terms]):
        p = row.cells[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(str(v)).font.size = Pt(9)

    # ── Spacer ──
    row = table.add_row(); row.cells[0].merge(row.cells[6])

    # ── Line item headers ──
    row = table.add_row()
    for i, h in enumerate(["ITEM#", "DESCRIPTION", "UNIT SIZE (in)", "AREA EACH (ft²)",
                            "QTY", "TOTAL AREA (ft²)", "TOTAL"]):
        row.cells[i].width = col_widths[i]
        p = row.cells[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h); run.bold = True; run.font.size = Pt(8); run.font.small_caps = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _add_shading(row.cells[i], '2E75B6')

    # ── Line item rows ──
    for idx, line in enumerate(glass_lines):
        row = table.add_row()
        line_total = line['area_total'] * price_per_sqft
        vals = [str(idx + 1), line.get('description', ''), line.get('size_str', ''),
                f"{line['area_each']:.2f}", str(line['qty']),
                f"{line['area_total']:.2f}", f"${line_total:,.2f}"]
        for i, v in enumerate(vals):
            row.cells[i].width = col_widths[i]
            p = row.cells[i].paragraphs[0]
            if i == 0 or i >= 3:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if i == 6:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.add_run(v).font.size = Pt(9)

    # ── Packaging + Subtotal ──
    row = table.add_row()
    c = row.cells[0]; c.merge(row.cells[3])
    run = c.paragraphs[0].add_run(f"Packaging: {packaging_note}"); run.underline = True; run.font.size = Pt(9)
    c = row.cells[4]; c.merge(row.cells[5])
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    c.paragraphs[0].add_run("SUBTOTAL").font.size = Pt(9)
    row.cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = row.cells[6].paragraphs[0].add_run(f"${subtotal:,.2f}"); run.bold = True; run.font.size = Pt(9)

    # ── Helper for cost rows ──
    def _cost_row(label, amount):
        r = table.add_row()
        r.cells[0].merge(r.cells[3])
        c2 = r.cells[4]; c2.merge(r.cells[5])
        c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        c2.paragraphs[0].add_run(label).font.size = Pt(9)
        r.cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r.cells[6].paragraphs[0].add_run(f"${amount:,.2f}" if amount else "").font.size = Pt(9)
        return r

    _cost_row("SALES TAX", sales_tax)
    _cost_row("PACKAGING", packaging_cost)

    # ── Shipping row (with terms text on the left) ──
    row = table.add_row()
    c = row.cells[0]; c.merge(row.cells[3])
    for ln in [
        "1. Enter this order in accordance with the prices, terms, delivery method, and specifications listed in this purchase order.",
        "2. Please notify us immediately if you are unable to ship as specified.",
        "3. Send all correspondence to:",
        "   INOVUES, INC.",
        "   2700 Post Oak Blvd, 2100, Houston, TX 77056",
        "   (833) 466-8837 (INO-VUES)",
        "   accounts@inovues.com",
    ]:
        c.add_paragraph().add_run(ln).font.size = Pt(7)
    c2 = row.cells[4]; c2.merge(row.cells[5])
    c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    c2.paragraphs[0].add_run("SHIPPING & HANDLING").font.size = Pt(9)
    row.cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    row.cells[6].paragraphs[0].add_run(f"${shipping_cost:,.2f}" if shipping_cost else "").font.size = Pt(9)

    _cost_row("OTHER", other_cost)

    # ── Total ──
    row = table.add_row()
    row.cells[0].merge(row.cells[3])
    c = row.cells[4]; c.merge(row.cells[5])
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = c.paragraphs[0].add_run("TOTAL"); run.bold = True; run.font.size = Pt(10)
    row.cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = row.cells[6].paragraphs[0].add_run(f"${total:,.2f}"); run.bold = True; run.font.size = Pt(10)

    # ── Signature ──
    row = table.add_row()
    row.cells[0].merge(row.cells[3])
    c = row.cells[4]; c.merge(row.cells[5])
    run = c.paragraphs[0].add_run("Authorized by _____________________"); run.italic = True; run.font.size = Pt(9)
    row.cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    row.cells[6].paragraphs[0].add_run(po_date).font.size = Pt(9)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# Display logo and title
st.image("ilogo.png", width=200)
st.title("SWR Cutlist")

# Project details inputs
project_name = st.text_input("Enter Project Name")
project_number = st.text_input("Enter Project Number", value="INO-")
prepared_by = st.text_input("Prepared By")

# System type and finish
system_type = st.selectbox("Select System Type", ["SWR-IG", "SWR-VIG", "SWR", "Custom"])
finish = st.selectbox("Select Finish", ["Mil Finish", "Clear Anodized", "Black Anodized", "Painted", "Light Bronze", "Medium Bronze", "Dark Bronze", "Champagne Anodized", "Gold Anodized", "Dark Gold Anodized"])

# Determine Glass Offset and profile number
offset_unit = 'mm'
if system_type == 'SWR-IG':
    glass_offset = 11.1125
    profile_number = '03003'
elif system_type == 'SWR-VIG':
    glass_offset = 11.1125
    profile_number = '03004'
elif system_type == 'SWR':
    swr_profile_choice = st.radio(
        "Select SWR Profile (Profile – Glass Offset)",
        [
            "03002 – 7.571 mm glass offset",
            "03111 – 11.11 mm glass offset",
        ]
    )
    if swr_profile_choice.startswith("03002"):
        profile_number = '03002'
        glass_offset = 7.571
    else:
        profile_number = '03111'
        glass_offset = 11.11
else:
    unit_offset = st.radio("Select Unit for Glass Offset", ["Inches", "Millimeters"], index=0)
    if unit_offset == 'Inches':
        glass_offset = st.number_input("Enter Glass Offset (in inches)", value=0.0) * inches_to_mm
        offset_unit = 'inches'
    else:
        glass_offset = st.number_input("Enter Glass Offset (in mm)", value=0.0)
        offset_unit = 'mm'

if system_type != 'Custom':
    st.write(f"Using a Glass Offset of {glass_offset:.3f} mm for system {system_type}")

# Input parameters and units
unit_tolerance = st.radio("Unit for Glass Cutting Tolerance", ["Inches", "Millimeters"], index=0)
tol_unit = 'inches' if unit_tolerance == 'Inches' else 'mm'
if unit_tolerance == 'Inches':
    glass_cutting_tolerance = st.number_input("Enter Glass Cutting Tolerance (in inches)", value=0.0625, format="%.4f")
else:
    val_mm = st.number_input("Enter Glass Cutting Tolerance (in mm)", value=0.0625 * inches_to_mm, format="%.3f")
    glass_cutting_tolerance = val_mm * mm_to_inches

unit_joint_top = st.radio("Unit for Joint Top", ["Inches", "Millimeters"], index=0)
top_unit = 'inches' if unit_joint_top == 'Inches' else 'mm'
if unit_joint_top == 'Inches':
    joint_top = st.number_input("Enter Joint Top (in inches)", value=0.5, format="%.3f")
else:
    val_mm = st.number_input("Enter Joint Top (in mm)", value=0.5 * inches_to_mm, format="%.3f")
    joint_top = val_mm * mm_to_inches

unit_joint_bottom = st.radio("Unit for Joint Bottom", ["Inches", "Millimeters"], index=0)
bottom_unit = 'inches' if unit_joint_bottom == 'Inches' else 'mm'
if unit_joint_bottom == 'Inches':
    joint_bottom = st.number_input("Enter Joint Bottom (in inches)", value=0.125, format="%.3f")
else:
    val_mm = st.number_input("Enter Joint Bottom (in mm)", value=0.125 * inches_to_mm, format="%.3f")
    joint_bottom = val_mm * mm_to_inches

unit_joint_left = st.radio("Unit for Joint Left", ["Inches", "Millimeters"], index=0)
left_unit = 'inches' if unit_joint_left == 'Inches' else 'mm'
if unit_joint_left == 'Inches':
    joint_left = st.number_input("Enter Joint Left (in inches)", value=0.25, format="%.3f")
else:
    val_mm = st.number_input("Enter Joint Left (in mm)", value=0.25 * inches_to_mm, format="%.3f")
    joint_left = val_mm * mm_to_inches

unit_joint_right = st.radio("Unit for Joint Right", ["Inches", "Millimeters"], index=0)
right_unit = 'inches' if unit_joint_right == 'Inches' else 'mm'
if unit_joint_right == 'Inches':
    joint_right = st.number_input("Enter Joint Right (in inches)", value=0.25, format="%.3f")
else:
    val_mm = st.number_input("Enter Joint Right (in mm)", value=0.25 * inches_to_mm, format="%.3f")
    joint_right = val_mm * mm_to_inches

# Determine part number
part_number = f"{system_type}-{profile_number}" if system_type != 'Custom' else 'Custom'

# Download template
template_path = 'SWR template.csv'
with open(template_path, 'rb') as f:
    st.download_button('Download Template', data=f.read(), file_name='SWR_template.csv', mime='text/csv')

# File upload
uploaded_file = st.file_uploader('Upload a CSV file', type='csv')

# Prepare parameters rows
params = [
    ('Prepared By', prepared_by, ''),
    ('System Type', system_type, ''),
    ('Finish', finish, ''),
    ('Glass Offset', glass_offset, offset_unit),
    ('Glass Cutting Tolerance', glass_cutting_tolerance, tol_unit),
    ('Joint Top', joint_top, top_unit),
    ('Joint Bottom', joint_bottom, bottom_unit),
    ('Joint Left', joint_left, left_unit),
    ('Joint Right', joint_right, right_unit)
]

# Process when file uploaded
if uploaded_file:
    df = pd.read_csv(uploaded_file)
    st.dataframe(df)

    # ── Date + version for filenames ──
    file_date = datetime.now().strftime("%Y-%m-%d")
    if 'file_version_date' not in st.session_state or st.session_state.file_version_date != file_date:
        st.session_state.file_version_date = file_date
        st.session_state.file_version = 1
    file_ver = st.session_state.file_version

    def make_fname(label, ext="xlsx"):
        """Generate filename: INO_{project}_{label}_{date}_v{ver}.{ext}"""
        return f"INO_{project_number}_{label}_{file_date}_v{file_ver}.{ext}"

    # Convert dims
    df['Overall Width mm'] = df['Overall Width in'] * inches_to_mm
    df['Overall Height mm'] = df['Overall Height in'] * inches_to_mm

    j_l = joint_left * inches_to_mm
    j_r = joint_right * inches_to_mm
    j_t = joint_top * inches_to_mm
    j_b = joint_bottom * inches_to_mm

    # SWR dims
    df['SWR Width mm'] = df['Overall Width mm'] - j_l - j_r
    df['SWR Height mm'] = df['Overall Height mm'] - j_t - j_b
    df['SWR Width in'] = df['SWR Width mm'] * mm_to_inches
    df['SWR Height in'] = df['SWR Height mm'] * mm_to_inches

    # Glass dims
    df['Glass Width mm'] = df['SWR Width mm'] - 2 * glass_offset
    df['Glass Height mm'] = df['SWR Height mm'] - 2 * glass_offset
    df['Glass Width in'] = df['Glass Width mm'] * mm_to_inches
    df['Glass Height in'] = df['Glass Height mm'] * mm_to_inches

    # Helper: sixteenth rounding
    def to_sixteenth(x):
        total = round(x * 16)
        w, f = divmod(total, 16)
        return f"{w} {f}/16" if f else f"{w}"

    # --- Glass File Export ---
    glass_df = pd.DataFrame({
        'Tag': df['Tag'].values,
        'Glass Width in': df['Glass Width in'].values,
        'Glass Width (1/16)': df['Glass Width in'].apply(to_sixteenth).values,
        'Glass Height in': df['Glass Height in'].values,
        'Glass Height (1/16)': df['Glass Height in'].apply(to_sixteenth).values,
        'Area Each (ft²)': (df['Glass Width in'] * df['Glass Height in']).values * sq_inches_to_sq_feet,
        'Qty': df['Qty'].values,
        'Area Total (ft²)': (df['Qty'] * (df['Glass Width in'] * df['Glass Height in']) * sq_inches_to_sq_feet).values
    })

    # Add totals row - FIX: only include numeric columns to avoid FutureWarning
    totals_data = {
        'Tag': 'Totals',
        'Glass Width in': '',
        'Glass Width (1/16)': '',
        'Glass Height in': '',
        'Glass Height (1/16)': '',
        'Area Each (ft²)': '',
        'Qty': glass_df['Qty'].sum(),
        'Area Total (ft²)': glass_df['Area Total (ft²)'].sum()
    }
    totals = pd.DataFrame([totals_data])
    glass_df = pd.concat([glass_df, totals], ignore_index=True)

    buf = BytesIO()
    with pd.ExcelWriter(buf, engine='xlsxwriter') as writer:
        ws = writer.book.add_worksheet('Glass')
        ws.insert_image('A1', 'ilogo.png', {'x_scale': 0.2, 'y_scale': 0.2})
        ws.write_row('A7', ['Project Name:', project_name])
        ws.write_row('A8', ['Project Number:', project_number])
        ws.write_row('A9', ['Date Created:', datetime.now().strftime('%Y-%m-%d %H:%M:%S')])
        ws.write_row('A10', ['Prepared By:', prepared_by])
        glass_df.to_excel(writer, sheet_name='Glass', startrow=12, index=False)
        pws = writer.book.add_worksheet('Parameters')
        for idx, (lbl, val, unit) in enumerate(params, start=1):
            pws.write(idx - 1, 0, lbl)
            pws.write(idx - 1, 1, val)
            pws.write(idx - 1, 2, unit)
    st.download_button('Download Glass File', buf.getvalue(), file_name=make_fname('SWR_Glass'))

    # --- AggCutOnly File Export ---
    df['Qty x 2'] = df['Qty'] * 2
    wc = df.groupby('SWR Width in')['Qty'].sum().sort_values(ascending=False)
    hc = df.groupby('SWR Height in')['Qty'].sum().sort_values(ascending=False)
    dims = pd.Index(list(wc.index) + list(hc.index)).unique()
    agg = pd.DataFrame(0, index=dims, columns=['Part #', 'Miter'] + list(df['Tag'].unique()) + ['Total QTY'])
    agg['Part #'] = part_number
    agg['Miter'] = '**'
    for _, r in df.iterrows():
        w, h, t, q2 = r['SWR Width in'], r['SWR Height in'], r['Tag'], r['Qty x 2']
        agg.at[w, t] += q2
        agg.at[h, t] += q2
    agg['Total QTY'] = agg[list(df['Tag'].unique())].sum(axis=1)
    agg = agg.reset_index().rename(columns={'index': 'Finished Length in'})

    buf2 = BytesIO()
    with pd.ExcelWriter(buf2, engine='xlsxwriter') as writer:
        ws = writer.book.add_worksheet('AggCutOnly')
        ws.insert_image('A1', 'ilogo.png', {'x_scale': 0.2, 'y_scale': 0.2})
        ws.write_row('A7', ['Project Name:', project_name])
        ws.write_row('A8', ['Project Number:', project_number])
        ws.write_row('A9', ['Date Created:', datetime.now().strftime('%Y-%m-%d %H:%M:%S')])
        ws.write_row('A10', ['Prepared By:', prepared_by])
        ws.write_row('A11', ['Finish:', finish])
        agg.to_excel(writer, sheet_name='AggCutOnly', startrow=12, index=False)
        pws = writer.book.add_worksheet('Parameters')
        for idx, (lbl, val, unit) in enumerate(params, start=1):
            pws.write(idx - 1, 0, lbl)
            pws.write(idx - 1, 1, val)
            pws.write(idx - 1, 2, unit)
    st.download_button('Download AggCutOnly File', buf2.getvalue(), file_name=make_fname('SWR_AggCutOnly'))

    # --- TagDetails File Export ---
    buf3 = BytesIO()
    with pd.ExcelWriter(buf3, engine='xlsxwriter') as writer:
        for tag in df['Tag'].unique():
            rows = {'Item': [], 'Position': [], 'Quantity': [], 'Length (mm)': [], 'Length (in)': []}
            subset = df[df['Tag'] == tag]
            for idx, r in subset.iterrows():
                for pos, length in [('left', r['SWR Width mm']), ('right', r['SWR Width mm']),
                                    ('top', r['SWR Height mm']), ('bottom', r['SWR Height mm'])]:
                    rows['Item'].append(idx + 1)
                    rows['Position'].append(pos)
                    rows['Quantity'].append(r['Qty'] * 2)
                    rows['Length (mm)'].append(length)
                    rows['Length (in)'].append(length * mm_to_inches)
            tag_df = pd.DataFrame(rows)
            ws = writer.book.add_worksheet(str(tag))
            ws.insert_image('A1', 'ilogo.png', {'x_scale': 0.2, 'y_scale': 0.2})
            ws.write_row('A7', ['Project Name:', project_name])
            ws.write_row('A8', ['Project Number:', project_number])
            ws.write_row('A9', ['Date Created:', datetime.now().strftime('%Y-%m-%d %H:%M:%S')])
            ws.write_row('A10', ['Prepared By:', prepared_by])
            tag_df.to_excel(writer, sheet_name=str(tag), startrow=12, index=False)
        pws = writer.book.add_worksheet('Parameters')
        for idx, (lbl, val, unit) in enumerate(params, start=1):
            pws.write(idx - 1, 0, lbl)
            pws.write(idx - 1, 1, val)
            pws.write(idx - 1, 2, unit)
    st.download_button('Download TagDetails File', buf3.getvalue(), file_name=make_fname('SWR_TagDetails'))

    # --- SWR Table File Export ---
    buf4 = BytesIO()
    with pd.ExcelWriter(buf4, engine='xlsxwriter') as writer:
        ws = writer.book.add_worksheet('Table')
        ws.insert_image('A1', 'ilogo.png', {'x_scale': 0.2, 'y_scale': 0.2})
        ws.write_row('A7', ['Project Name:', project_name])
        ws.write_row('A8', ['Project Number:', project_number])
        ws.write_row('A9', ['Date Created:', datetime.now().strftime('%Y-%m-%d %H:%M:%S')])
        ws.write_row('A10', ['Prepared By:', prepared_by])
        ws.write_row('A11', ['Finish:', finish])
        ws.write_row('A12', ['Glass Cutting Tolerance:', glass_cutting_tolerance])
        ws.write_row('A13', ['Tolerance Unit:', tol_unit])
        ws.write_row('A14', ['Joint Top:', joint_top])
        ws.write_row('A15', ['Joint Top Unit:', top_unit])
        ws.write_row('A16', ['Joint Bottom:', joint_bottom])
        ws.write_row('A17', ['Joint Bottom Unit:', bottom_unit])
        ws.write_row('A18', ['Joint Left:', joint_left])
        ws.write_row('A19', ['Joint Left Unit:', left_unit])
        ws.write_row('A20', ['Joint Right:', joint_right])
        ws.write_row('A21', ['Joint Right Unit:', right_unit])
        df.drop(columns=['Qty x 2'], errors='ignore').to_excel(writer, sheet_name='Table', startrow=23, index=False)
        pws = writer.book.add_worksheet('Parameters')
        for idx, (lbl, val, unit) in enumerate(params, start=1):
            pws.write(idx - 1, 0, lbl)
            pws.write(idx - 1, 1, val)
            pws.write(idx - 1, 2, unit)
    st.download_button('Download SWR Table File', buf4.getvalue(), file_name=make_fname('SWR_Table'))

    # ─────────────────────────────────────────────────────────────
    # 💾 SAVE TO ODOO
    # ─────────────────────────────────────────────────────────────
    st.divider()
    st.subheader("💾 Save to Odoo Project")

    ODOO_URL     = os.environ.get("ODOO_URL",     "https://inovues.odoo.com")
    ODOO_DB      = os.environ.get("ODOO_DB",      "inovues")
    ODOO_USER    = os.environ.get("ODOO_USER",    "sketterer@inovues.com")
    ODOO_API_KEY = os.environ.get("ODOO_API_KEY", "")

    @st.cache_data(ttl=300, show_spinner="Loading projects from Odoo...")
    def fetch_odoo_projects():
        try:
            common = xmlrpc.client.ServerProxy(f"{ODOO_URL}/xmlrpc/2/common")
            uid    = common.authenticate(ODOO_DB, ODOO_USER, ODOO_API_KEY, {})
            if not uid:
                return None, "Authentication failed — check ODOO_API_KEY env var."
            models = xmlrpc.client.ServerProxy(f"{ODOO_URL}/xmlrpc/2/object")
            projects = models.execute_kw(
                ODOO_DB, uid, ODOO_API_KEY,
                "project.project", "search_read",
                [[["active", "=", True]]],
                {"fields": ["id", "name"], "order": "name asc", "limit": 200}
            )
            return {p["name"]: p["id"] for p in projects}, None
        except Exception as e:
            return None, str(e)

    project_map, err = fetch_odoo_projects()

    if err:
        st.error(f"❌ Could not load Odoo projects: {err}")
    elif not project_map:
        st.warning("No active projects found in Odoo.")
    else:
        selected_project_name = st.selectbox(
            "Select Odoo Project to attach files to:",
            options=list(project_map.keys()),
            index=None,
            placeholder="Choose a project..."
        )

        if st.button("📎 Attach all 4 files to selected project", type="primary",
                     disabled=selected_project_name is None):
            project_id = project_map[selected_project_name]
            with st.spinner(f"Attaching files to '{selected_project_name}'..."):
                try:
                    common = xmlrpc.client.ServerProxy(f"{ODOO_URL}/xmlrpc/2/common")
                    uid    = common.authenticate(ODOO_DB, ODOO_USER, ODOO_API_KEY, {})
                    models = xmlrpc.client.ServerProxy(f"{ODOO_URL}/xmlrpc/2/object")

                    def odoo_call(model, method, args, kwargs={}):
                        return models.execute_kw(ODOO_DB, uid, ODOO_API_KEY, model, method, args, kwargs)

                    ts = datetime.now().strftime("%Y%m%d_%H%M%S")

                    # Create a task in the Engineering stage, set to Approved
                    task_id = odoo_call("project.task", "create", [{
                        "name": f"SWR Cutlist — {datetime.now().strftime('%Y-%m-%d')}",
                        "project_id": project_id,
                        "stage_id": 8,  # Engineering
                        "state": "03_approved",
                    }])

                    files = [
                        (make_fname('SWR_Glass'),      buf.getvalue()),
                        (make_fname('SWR_AggCutOnly'),  buf2.getvalue()),
                        (make_fname('SWR_TagDetails'),  buf3.getvalue()),
                        (make_fname('SWR_Table'),       buf4.getvalue()),
                    ]

                    for fname, fdata in files:
                        odoo_call("ir.attachment", "create", [{
                            "name":      fname,
                            "type":      "binary",
                            "datas":     base64.b64encode(fdata).decode("utf-8"),
                            "res_model": "project.task",
                            "res_id":    task_id,
                            "mimetype":  "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        }])

                    odoo_call("project.task", "message_post", [[task_id]], {
                        "body": (
                            f"<b>✂️ SWR Cut List attached</b><br/>"
                            f"Prepared by: {prepared_by}<br/>"
                            f"System: {system_type} | Profile: {profile_number} | Finish: {finish}<br/>"
                            f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}<br/>"
                            f"Files: Glass, AggCutOnly, TagDetails, SWR Table"
                        ),
                        "message_type": "comment",
                        "subtype_xmlid": "mail.mt_comment",
                    })

                    st.success(f"✅ All 4 files (v{file_ver}) attached to task **SWR Cutlist** in **{selected_project_name}** (Engineering → Approved)!")
                    st.session_state.file_version += 1

                    # ── Also upload to Google Drive ──
                    if GDRIVE_WEBHOOK_URL:
                        gdrive_ok = 0
                        gdrive_subfolder = selected_project_name
                        for fname, fdata in files:
                            result, err = upload_to_gdrive(
                                fname, fdata,
                                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                subfolder_name=gdrive_subfolder
                            )
                            if result:
                                gdrive_ok += 1
                            elif err:
                                st.warning(f"⚠️ Drive upload failed for {fname}: {err}")
                        if gdrive_ok > 0:
                            st.success(f"📁 {gdrive_ok} file(s) also saved to Google Drive → {gdrive_subfolder}/")

                except xmlrpc.client.Fault as e:
                    st.error(f"❌ Odoo API error: {e.faultString}")
                except Exception as e:
                    st.error(f"❌ Error: {str(e)}")

    # ─────────────────────────────────────────────────────────────
    # 📝 CREATE PURCHASE ORDER (DOCX + ODOO)
    # ─────────────────────────────────────────────────────────────
    st.divider()
    st.subheader("📝 Glass Purchase Order")

    glass_lines = glass_df[glass_df['Tag'] != 'Totals'].copy()

    total_qty = int(pd.to_numeric(glass_lines['Qty']).sum())
    total_area = pd.to_numeric(glass_lines['Area Total (ft²)']).sum()

    # --- Per-line description inputs ---
    st.write("**PO Line Items — enter description per size (first entry becomes default for all):**")

    # Build size list for display
    size_strs = []
    for _, row in glass_lines.iterrows():
        size_strs.append(f"{row['Glass Width (1/16)']}\" x {row['Glass Height (1/16)']}\"")

    # First line description sets the default
    first_desc = st.text_input(
        f"Description for all sizes (or just line 1: {size_strs[0]})",
        value="",
        key="po_desc_0",
        placeholder="e.g. GT1 – 10mm Leadus VIG: 5Tlow-E(D80)+V+5T"
    )

    line_descriptions = [first_desc]
    if len(size_strs) > 1:
        with st.expander(f"Override descriptions for lines 2–{len(size_strs)} (default: same as line 1)", expanded=False):
            for i in range(1, len(size_strs)):
                desc = st.text_input(
                    f"Line {i+1}: {size_strs[i]}",
                    value="",
                    key=f"po_desc_{i}",
                    placeholder=first_desc if first_desc else "Same as line 1"
                )
                line_descriptions.append(desc if desc else first_desc)

    # Preview table
    preview_data = []
    for i, (_, row) in enumerate(glass_lines.iterrows()):
        desc = line_descriptions[i] if i < len(line_descriptions) else first_desc
        preview_data.append({
            'Description': desc,
            'Size': size_strs[i],
            'Area Each (ft²)': round(float(row['Area Each (ft²)']), 2),
            'Qty': int(float(row['Qty'])),
            'Area Total (ft²)': round(float(row['Area Total (ft²)']), 2),
        })
    st.dataframe(pd.DataFrame(preview_data), use_container_width=True, hide_index=True)
    st.write(f"**Totals:** {total_qty} pieces | {total_area:.2f} ft²")

    # --- Vendor selector (with full contact details) ---
    @st.cache_data(ttl=300, show_spinner="Loading vendors from Odoo...")
    def fetch_odoo_vendors():
        try:
            _common = xmlrpc.client.ServerProxy(f"{ODOO_URL}/xmlrpc/2/common")
            _uid = _common.authenticate(ODOO_DB, ODOO_USER, ODOO_API_KEY, {})
            if not _uid:
                return None, "Authentication failed."
            _models = xmlrpc.client.ServerProxy(f"{ODOO_URL}/xmlrpc/2/object")
            vendor_fields = ["id", "name", "email", "phone", "street", "street2",
                             "city", "state_id", "zip", "country_id",
                             "child_ids"]
            vendors = _models.execute_kw(
                ODOO_DB, _uid, ODOO_API_KEY,
                "res.partner", "search_read",
                [[["supplier_rank", ">", 0]]],
                {"fields": vendor_fields, "order": "name asc", "limit": 200}
            )
            if not vendors:
                vendors = _models.execute_kw(
                    ODOO_DB, _uid, ODOO_API_KEY,
                    "res.partner", "search_read",
                    [[["is_company", "=", True]]],
                    {"fields": vendor_fields, "order": "name asc", "limit": 200}
                )
            # Build contact name: prefer child tagged 'Orders', else first child.
            # The 'Orders' tag (res.partner.category) flags the primary ordering
            # contact for a vendor — set it on whichever child should appear in POs.
            for v in vendors:
                if v.get("child_ids"):
                    # Try to find a child tagged 'Orders' first
                    primary_ids = _models.execute_kw(
                        ODOO_DB, _uid, ODOO_API_KEY,
                        "res.partner", "search",
                        [[
                            ["id", "in", v["child_ids"]],
                            ["category_id.name", "=", "Orders"],
                        ]],
                        {"limit": 1}
                    )
                    target_ids = primary_ids if primary_ids else v["child_ids"][:1]
                    contacts = _models.execute_kw(
                        ODOO_DB, _uid, ODOO_API_KEY,
                        "res.partner", "read",
                        [target_ids],
                        {"fields": ["name", "email"]}
                    )
                    if contacts:
                        v["contact_name"] = contacts[0].get("name", "")
                        if not v.get("email"):
                            v["email"] = contacts[0].get("email", "")
                else:
                    v["contact_name"] = ""
                # Build address string
                addr_parts = [p for p in [v.get("street"), v.get("street2")] if p]
                city_line = ", ".join(p for p in [
                    v.get("city"),
                    v.get("state_id", [False, ""])[1] if isinstance(v.get("state_id"), list) else "",
                    v.get("zip")
                ] if p)
                if city_line:
                    addr_parts.append(city_line)
                country = v.get("country_id", [False, ""])[1] if isinstance(v.get("country_id"), list) else ""
                if country:
                    addr_parts.append(country)
                v["full_address"] = "\n".join(addr_parts)

            return {v["name"]: v for v in vendors}, None
        except Exception as e:
            return None, str(e)

    vendor_map, vendor_err = fetch_odoo_vendors()
    selected_vendor = None

    if vendor_err:
        st.error(f"❌ Could not load vendors: {vendor_err}")
    elif not vendor_map:
        st.warning("No vendors found in Odoo.")
    else:
        selected_vendor = st.selectbox(
            "Select Glass Vendor",
            options=list(vendor_map.keys()),
            index=None,
            placeholder="Choose a vendor...",
            help="Required — vendor details will be pulled from Odoo."
        )

    # Show vendor details if selected
    if selected_vendor and vendor_map:
        v = vendor_map[selected_vendor]
        with st.expander("📇 Vendor Details (from Odoo)", expanded=False):
            st.text(f"Name: {v['name']}")
            st.text(f"Contact: {v.get('contact_name', '')}")
            st.text(f"Email: {v.get('email', '')}")
            st.text(f"Address: {v.get('full_address', '')}")

    st.divider()

    # --- PO Details & Pricing ---
    st.write("**PO Details**")
    po_col1, po_col2, po_col3 = st.columns(3)
    with po_col1:
        po_number_input = st.text_input("PO Number", value=project_number)
        po_requisitioner = st.text_input("Requisitioner", value=prepared_by)
    with po_col2:
        po_lead_time = st.text_input("Lead Time", value="ASAP")
        po_shipped_via = st.selectbox("Shipped Via", options=["Air", "Ground"], index=0)
    with po_col3:
        po_fob = st.selectbox("F.O.B. Point", options=["CIF", "DDP", "DAP", "EXW", "FOB"], index=1)
        po_terms = st.text_input("Terms", value="Net 30")

    st.write("**Ship To**")
    ship_to_default = "Momentum Glass, LLC\nAttn: INOVUES, INC.\n25825 Aldine Westfield Rd.\nSpring, TX 77373\n281.809.2830"
    ship_to_text = st.text_area("Ship To Address", value=ship_to_default, height=120)
    job_location = st.text_input("Job Location", value="")

    st.write("**Pricing**")
    pr_col1, pr_col2 = st.columns(2)
    with pr_col1:
        price_per_sqft = st.number_input("Price per sqft ($)", value=0.0, min_value=0.0, step=0.01, format="%.2f")
        packaging_cost = st.number_input("Packaging ($)", value=0.0, min_value=0.0, step=10.0, format="%.2f")
        sales_tax = st.number_input("Sales Tax ($)", value=0.0, min_value=0.0, step=0.01, format="%.2f")
    with pr_col2:
        shipping_cost = st.number_input("Shipping & Handling ($)", value=0.0, min_value=0.0, step=50.0, format="%.2f")
        other_cost = st.number_input("Other ($)", value=0.0, min_value=0.0, step=10.0, format="%.2f")
        packaging_note = st.text_input("Packaging Note", value="Non-returnable boxed crate/rack")

    # Calculate and show totals
    subtotal = total_area * price_per_sqft
    grand_total = subtotal + sales_tax + packaging_cost + shipping_cost + other_cost
    st.write(f"**Subtotal:** ${subtotal:,.2f} | **Grand Total:** ${grand_total:,.2f}")

    vendor_ready = (vendor_map and not vendor_err and selected_vendor is not None)

    # --- Build glass line items for the docx ---
    po_glass_lines = []
    for i, (_, row) in enumerate(glass_lines.iterrows()):
        desc = line_descriptions[i] if i < len(line_descriptions) else first_desc
        po_glass_lines.append({
            'description': desc,
            'size_str': size_strs[i],
            'area_each': round(float(row['Area Each (ft²)']), 2),
            'qty': int(float(row['Qty'])),
            'area_total': round(float(row['Area Total (ft²)']), 2),
        })

    st.divider()

    # --- Generate DOCX PO ---
    doc_col, odoo_col = st.columns(2)

    with doc_col:
        st.write("**📄 Download PO Document**")
        if st.button("📄 Generate PO (.docx)", type="secondary", disabled=not vendor_ready):
            v = vendor_map[selected_vendor]
            logo_path = "ilogo.png" if os.path.exists("ilogo.png") else None
            po_buf = generate_po_docx(
                vendor_name=v['name'],
                vendor_contact=v.get('contact_name', ''),
                vendor_email=v.get('email', ''),
                vendor_address=v.get('full_address', ''),
                ship_to_lines=[ln for ln in ship_to_text.split('\n') if ln.strip()],
                job_number=project_number,
                job_location=job_location,
                po_date=datetime.now().strftime('%m/%d/%Y'),
                po_number=po_number_input,
                requisitioner=po_requisitioner,
                lead_time=po_lead_time,
                shipped_via=po_shipped_via,
                fob_point=po_fob,
                terms=po_terms,
                glass_lines=po_glass_lines,
                price_per_sqft=price_per_sqft,
                packaging_cost=packaging_cost,
                shipping_cost=shipping_cost,
                sales_tax=sales_tax,
                other_cost=other_cost,
                packaging_note=packaging_note,
                logo_path=logo_path,
            )
            st.session_state['po_docx_buf'] = po_buf.getvalue()
            st.success("✅ PO document generated!")

        if 'po_docx_buf' in st.session_state:
            st.download_button(
                "💾 Download PO .docx",
                data=st.session_state['po_docx_buf'],
                file_name=make_fname(f"PO_{po_number_input}", "docx"),
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

    # --- Create Draft PO in Odoo + attach docx ---
    with odoo_col:
        st.write("**🔗 Create Draft PO in Odoo**")
        if st.button("📝 Create Draft PO in Odoo", type="primary", disabled=not vendor_ready):
            with st.spinner("Creating draft Purchase Order in Odoo..."):
                try:
                    common = xmlrpc.client.ServerProxy(f"{ODOO_URL}/xmlrpc/2/common")
                    uid    = common.authenticate(ODOO_DB, ODOO_USER, ODOO_API_KEY, {})
                    models = xmlrpc.client.ServerProxy(f"{ODOO_URL}/xmlrpc/2/object")

                    def odoo_call(model, method, args, kwargs={}):
                        return models.execute_kw(ODOO_DB, uid, ODOO_API_KEY, model, method, args, kwargs)

                    product_ids = odoo_call("product.product", "search",
                        [[["name", "ilike", "SWR Glass Panel"]]])

                    if not product_ids:
                        st.error("❌ Product 'SWR Glass Panel' not found in Odoo. Please create it first.")
                    else:
                        product_id = product_ids[0]
                        v = vendor_map[selected_vendor]

                        order_lines = []
                        for gl in po_glass_lines:
                            odoo_desc = (
                                f"{gl['description']}\n"
                                f"{gl['size_str']}\n"
                                f"Unit Area: {gl['area_each']} ft2  |  Qty: {gl['qty']} pcs  |  Total Area: {gl['area_total']} ft2"
                            )
                            line_total = gl['area_total'] * price_per_sqft
                            order_lines.append((0, 0, {
                                "product_id":  product_id,
                                "name":        odoo_desc,
                                "product_qty": gl['qty'],
                                "price_unit":  line_total / gl['qty'] if gl['qty'] > 0 else 0.0,
                            }))

                        po_vals = {
                            "partner_id": v["id"],
                            "origin": f"INO-{project_number} / SWR Cut List",
                            "order_line": order_lines,
                        }

                        po_id = odoo_call("purchase.order", "create", [po_vals])

                        po_data = odoo_call("purchase.order", "read",
                            [po_id], {"fields": ["name"]})
                        po_name = po_data[0]["name"] if po_data else f"ID {po_id}"

                        # Generate and attach the PO docx to the Odoo task
                        logo_path = "ilogo.png" if os.path.exists("ilogo.png") else None
                        po_buf = generate_po_docx(
                            vendor_name=v['name'],
                            vendor_contact=v.get('contact_name', ''),
                            vendor_email=v.get('email', ''),
                            vendor_address=v.get('full_address', ''),
                            ship_to_lines=[ln for ln in ship_to_text.split('\n') if ln.strip()],
                            job_number=project_number,
                            job_location=job_location,
                            po_date=datetime.now().strftime('%m/%d/%Y'),
                            po_number=po_number_input,
                            requisitioner=po_requisitioner,
                            lead_time=po_lead_time,
                            shipped_via=po_shipped_via,
                            fob_point=po_fob,
                            terms=po_terms,
                            glass_lines=po_glass_lines,
                            price_per_sqft=price_per_sqft,
                            packaging_cost=packaging_cost,
                            shipping_cost=shipping_cost,
                            sales_tax=sales_tax,
                            other_cost=other_cost,
                            packaging_note=packaging_note,
                            logo_path=logo_path,
                        )

                        # Attach PO docx to the Odoo PO record
                        odoo_call("ir.attachment", "create", [{
                            "name":      make_fname(f"PO_{po_number_input}", "docx"),
                            "type":      "binary",
                            "datas":     base64.b64encode(po_buf.getvalue()).decode("utf-8"),
                            "res_model": "purchase.order",
                            "res_id":    po_id,
                            "mimetype":  "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        }])

                        # Post chatter message on the PO
                        odoo_call("purchase.order", "message_post", [[po_id]], {
                            "body": (
                                f"<b>📝 Draft PO from SWR Cutlist App</b><br/>"
                                f"Project: {project_name}<br/>"
                                f"Project Number: {project_number}<br/>"
                                f"System: {system_type} | Finish: {finish}<br/>"
                                f"Vendor: {v['name']}<br/>"
                                f"Prepared by: {prepared_by}<br/>"
                                f"Subtotal: ${subtotal:,.2f} | Total: ${grand_total:,.2f}<br/>"
                                f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}<br/>"
                                f"PO document (.docx) attached."
                            ),
                            "message_type": "comment",
                            "subtype_xmlid": "mail.mt_comment",
                        })

                        st.success(
                            f"✅ Draft PO **{po_name}** created for **{selected_vendor}** with "
                            f"{len(order_lines)} glass lines! PO docx attached.\n\n"
                            f"Open it in Odoo to review and confirm."
                        )
                        st.info(
                            f"🔗 Open in Odoo: {ODOO_URL}/web#id={po_id}"
                            f"&model=purchase.order&view_type=form"
                        )

                except xmlrpc.client.Fault as e:
                    st.error(f"❌ Odoo API error: {e.faultString}")
                except Exception as e:
                    st.error(f"❌ Error: {str(e)}")
